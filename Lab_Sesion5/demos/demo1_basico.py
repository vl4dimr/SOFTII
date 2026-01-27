"""
DEMO 1: Crear documento basico
Ejecutar: python demos/demo1_basico.py
"""
from docx import Document

# Crear documento vacio
doc = Document()

# Agregar titulo
doc.add_heading("Mi Primera Tesis", level=0)

# Agregar parrafos
doc.add_paragraph("Este es el contenido de mi tesis.")
doc.add_heading("Capitulo 1: Introduccion", level=1)
doc.add_paragraph("Aqui va la introduccion de la tesis.")

# Guardar
doc.save("demos/demo1_resultado.docx")
print("Documento creado: demos/demo1_resultado.docx")
print("Abrelo en Word o LibreOffice para ver el resultado.")
