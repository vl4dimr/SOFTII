"""
DEMO 2: Formato de texto con Runs
Ejecutar: python demos/demo2_formato.py
"""
from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading("Formato de Texto", level=0)

# Parrafo con multiples runs
p = doc.add_paragraph()
run1 = p.add_run("Este texto es normal, ")
run1.font.name = "Times New Roman"
run1.font.size = Pt(12)

run2 = p.add_run("este esta en negrita, ")
run2.bold = True
run2.font.name = "Times New Roman"
run2.font.size = Pt(12)

run3 = p.add_run("este en cursiva, ")
run3.italic = True
run3.font.name = "Times New Roman"
run3.font.size = Pt(12)

run4 = p.add_run("y este subrayado.")
run4.underline = True
run4.font.name = "Times New Roman"
run4.font.size = Pt(12)

# Otro parrafo con tamanio diferente
p2 = doc.add_paragraph()
run_grande = p2.add_run("Texto grande (24pt)")
run_grande.font.size = Pt(24)
run_grande.bold = True

doc.save("demos/demo2_resultado.docx")
print("Documento creado: demos/demo2_resultado.docx")
print("Observa los diferentes formatos en cada run.")
