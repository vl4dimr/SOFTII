"""
DEMO 3: Margenes formato UNAP
Ejecutar: python demos/demo3_margenes.py
"""
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
section = doc.sections[0]

# Margenes UNAP
section.left_margin = Cm(4)      # Izquierdo: 4cm (empaste)
section.right_margin = Cm(2.5)   # Derecho: 2.5cm
section.top_margin = Cm(2.5)     # Superior: 2.5cm
section.bottom_margin = Cm(2.5)  # Inferior: 2.5cm

# Contenido de prueba
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UNIVERSIDAD NACIONAL DEL ALTIPLANO")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(14)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("FACULTAD DE INGENIERIA ESTADISTICA E INFORMATICA")
run2.bold = True
run2.font.name = "Times New Roman"
run2.font.size = Pt(12)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run3 = p3.add_run(
    "Este parrafo tiene margenes configurados segun el formato UNAP. "
    "El margen izquierdo es de 4 centimetros para permitir el empaste. "
    "Los demas margenes son de 2.5 centimetros. La fuente es Times New Roman "
    "a 12 puntos con alineacion justificada."
)
run3.font.name = "Times New Roman"
run3.font.size = Pt(12)

doc.save("demos/demo3_resultado.docx")
print("Documento creado: demos/demo3_resultado.docx")
print("Abre en Word y activa la regla para ver los margenes.")
print("  Izquierdo: 4 cm | Derecho: 2.5 cm")
print("  Superior: 2.5 cm | Inferior: 2.5 cm")
