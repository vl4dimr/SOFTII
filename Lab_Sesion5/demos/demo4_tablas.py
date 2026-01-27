"""
DEMO 4: Crear tablas (jurados)
Ejecutar: python demos/demo4_tablas.py
"""
from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading("Jurado Calificador", level=1)

# Crear tabla de jurados
jurados = [
    ("PRESIDENTE", "Dr. Juan Carlos Perez Lopez"),
    ("PRIMER MIEMBRO", "Mg. Maria Elena Garcia Quispe"),
    ("SEGUNDO MIEMBRO", "Mg. Carlos Alberto Lopez Mamani"),
    ("DIRECTOR", "Dr. Roberto Flores Condori"),
]

tabla = doc.add_table(rows=len(jurados) + 1, cols=3)
tabla.style = "Table Grid"

# Encabezados
for i, texto in enumerate(["CARGO", "NOMBRE", "FIRMA"]):
    cell = tabla.rows[0].cells[i]
    cell.text = texto
    for run in cell.paragraphs[0].runs:
        run.bold = True

# Datos
for i, (cargo, nombre) in enumerate(jurados):
    tabla.rows[i + 1].cells[0].text = cargo
    tabla.rows[i + 1].cells[1].text = nombre
    tabla.rows[i + 1].cells[2].text = ""  # Espacio para firma

doc.save("demos/demo4_resultado.docx")
print("Documento creado: demos/demo4_resultado.docx")
print(f"Tabla con {len(jurados)} jurados creada.")
