"""
SERVICIO: GENERADOR DE DOCUMENTOS DOCX
Genera tesis en formato Word con formato UNAP

Sesion 5: Generacion de Documentos
Curso: Desarrollo de Software - UNAP
"""

from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def configurar_margenes_unap(doc: Document) -> None:
    """
    Configura los margenes segun formato UNAP.

    - Izquierdo: 4 cm (para empaste)
    - Derecho: 2.5 cm
    - Superior: 2.5 cm
    - Inferior: 2.5 cm
    """
    section = doc.sections[0]
    section.left_margin = Cm(4)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def agregar_parrafo_centrado(doc: Document, texto: str,
                             negrita: bool = False,
                             tamano: int = 12) -> None:
    """Agrega un parrafo centrado con formato opcional."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = negrita
    run.font.name = "Times New Roman"
    run.font.size = Pt(tamano)


def generar_portada(doc: Document, datos: dict) -> None:
    """Genera la portada de la tesis."""
    # Universidad
    agregar_parrafo_centrado(doc, "UNIVERSIDAD NACIONAL DEL ALTIPLANO",
                             negrita=True, tamano=14)
    agregar_parrafo_centrado(doc, "FACULTAD DE INGENIERIA ESTADISTICA E INFORMATICA",
                             negrita=True, tamano=12)
    agregar_parrafo_centrado(doc, datos.get("escuela", "ESCUELA PROFESIONAL DE INGENIERIA DE SISTEMAS"),
                             negrita=True, tamano=12)

    # Espacio
    doc.add_paragraph()
    doc.add_paragraph()

    # Titulo
    agregar_parrafo_centrado(doc, datos.get("titulo", "TITULO DE LA TESIS").upper(),
                             negrita=True, tamano=14)

    # Espacio
    doc.add_paragraph()

    # Tipo de trabajo
    agregar_parrafo_centrado(doc, "TESIS", negrita=True, tamano=12)
    agregar_parrafo_centrado(doc, "PRESENTADA POR:", tamano=12)
    agregar_parrafo_centrado(doc, datos.get("autor", "NOMBRE DEL AUTOR").upper(),
                             negrita=True, tamano=12)

    # Espacio
    doc.add_paragraph()

    # Para optar
    agregar_parrafo_centrado(doc, "PARA OPTAR EL TITULO PROFESIONAL DE:", tamano=12)
    agregar_parrafo_centrado(doc, datos.get("titulo_profesional", "INGENIERO DE SISTEMAS"),
                             negrita=True, tamano=12)

    # Espacio
    doc.add_paragraph()
    doc.add_paragraph()

    # Lugar y fecha
    agregar_parrafo_centrado(doc, "PUNO - PERU", tamano=12)
    agregar_parrafo_centrado(doc, datos.get("anio", "2026"), tamano=12)


def generar_tabla_jurados(doc: Document, jurados: list) -> None:
    """
    Genera la tabla de jurados.

    jurados: Lista de tuplas (cargo, nombre, grado)
    Ejemplo: [("PRESIDENTE", "Juan Perez", "Dr."), ...]
    """
    doc.add_page_break()

    agregar_parrafo_centrado(doc, "JURADO CALIFICADOR", negrita=True, tamano=12)
    doc.add_paragraph()

    # Crear tabla
    tabla = doc.add_table(rows=len(jurados) + 1, cols=3)
    tabla.style = "Table Grid"

    # Encabezados
    encabezados = tabla.rows[0].cells
    encabezados[0].text = "CARGO"
    encabezados[1].text = "NOMBRE"
    encabezados[2].text = "FIRMA"

    # Datos
    for i, (cargo, nombre, grado) in enumerate(jurados):
        fila = tabla.rows[i + 1].cells
        fila[0].text = cargo
        fila[1].text = f"{grado} {nombre}"
        fila[2].text = ""  # Espacio para firma


def generar_tesis(datos: dict) -> Document:
    """
    Genera un documento de tesis completo.

    Args:
        datos: Diccionario con:
            - titulo: Titulo de la tesis
            - autor: Nombre del autor
            - escuela: Escuela profesional (opcional)
            - titulo_profesional: Titulo a obtener (opcional)
            - anio: Anio de presentacion (opcional)
            - jurados: Lista de jurados (opcional)

    Returns:
        Documento Word listo para guardar
    """
    doc = Document()

    # Configurar margenes
    configurar_margenes_unap(doc)

    # Generar portada
    generar_portada(doc, datos)

    # Generar tabla de jurados si se proporcionan
    if "jurados" in datos and datos["jurados"]:
        generar_tabla_jurados(doc, datos["jurados"])

    return doc


def generar_desde_plantilla(plantilla_path: str, datos: dict) -> Document:
    """
    Genera documento reemplazando variables {{VAR}} en una plantilla.

    Args:
        plantilla_path: Ruta a la plantilla .docx
        datos: Diccionario con variables a reemplazar

    Returns:
        Documento con variables reemplazadas
    """
    doc = Document(plantilla_path)

    # Reemplazar en parrafos
    for parrafo in doc.paragraphs:
        for key, value in datos.items():
            marcador = "{{" + key + "}}"
            if marcador in parrafo.text:
                parrafo.text = parrafo.text.replace(marcador, str(value))

    # Reemplazar en tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for key, value in datos.items():
                    marcador = "{{" + key + "}}"
                    if marcador in celda.text:
                        celda.text = celda.text.replace(marcador, str(value))

    return doc
